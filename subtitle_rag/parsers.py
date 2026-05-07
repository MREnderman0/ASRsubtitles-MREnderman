from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd


TEXT_SUFFIXES = {".txt", ".md", ".srt", ".vtt"}
TABLE_SUFFIXES = {".csv", ".xlsx", ".xls"}
SUPPORTED_REFERENCE_SUFFIXES = TEXT_SUFFIXES | TABLE_SUFFIXES | {".pdf", ".docx"}
SUPPORTED_GLOSSARY_SUFFIXES = {".csv", ".xlsx", ".xls"}


@dataclass(frozen=True)
class GlossaryEntry:
    alias: str
    canonical: str
    note: str = ""
    source: str = ""


@dataclass(frozen=True)
class ReferenceDoc:
    name: str
    text: str


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _frame_to_text(df: pd.DataFrame) -> str:
    values: list[str] = []
    for value in df.fillna("").astype(str).to_numpy().ravel():
        value = value.strip()
        if value:
            values.append(value)
    return "\n".join(values)


def _read_table(path: Path) -> str:
    if path.suffix.lower() == ".csv":
        for encoding in ("utf-8-sig", "utf-8", "gbk"):
            try:
                return _frame_to_text(pd.read_csv(path, encoding=encoding))
            except UnicodeDecodeError:
                continue
        return _frame_to_text(pd.read_csv(path, encoding="utf-8", encoding_errors="ignore"))

    frames = pd.read_excel(path, sheet_name=None)
    return "\n".join(_frame_to_text(frame) for frame in frames.values())


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires `pip install -r subtitle_rag/requirements-extra.txt`.") from exc

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires `pip install -r subtitle_rag/requirements-extra.txt`.") from exc

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_cells: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)
    return "\n".join(paragraphs + table_cells)


def read_reference_file(path: str | Path) -> ReferenceDoc:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        text = _read_text(path)
    elif suffix in TABLE_SUFFIXES:
        text = _read_table(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    else:
        raise ValueError(f"Unsupported reference file type: {path.name}")
    return ReferenceDoc(name=path.name, text=text.strip())


def read_reference_files(paths: Iterable[str | Path]) -> list[ReferenceDoc]:
    docs: list[ReferenceDoc] = []
    for path in paths:
        doc = read_reference_file(path)
        if doc.text:
            docs.append(doc)
    return docs


def _read_glossary_frame(path: Path) -> pd.DataFrame:
    if path.suffix.lower() == ".csv":
        for encoding in ("utf-8-sig", "utf-8", "gbk"):
            try:
                return pd.read_csv(path, encoding=encoding, header=None)
            except UnicodeDecodeError:
                continue
        return pd.read_csv(path, encoding="utf-8", encoding_errors="ignore", header=None)
    return pd.read_excel(path, header=None)


def read_glossary_files(paths: Iterable[str | Path]) -> list[GlossaryEntry]:
    entries: list[GlossaryEntry] = []
    for raw_path in paths:
        path = Path(raw_path)
        if path.suffix.lower() not in SUPPORTED_GLOSSARY_SUFFIXES:
            continue
        frame = _read_glossary_frame(path).fillna("")
        if frame.empty:
            continue

        start_row, alias_col, canonical_col, note_col = _glossary_columns(frame)

        for _, row in frame.iloc[start_row:].iterrows():
            alias = str(row.get(alias_col, "")).strip()
            canonical = str(row.get(canonical_col, "")).strip() or alias
            note = str(row.get(note_col, "")).strip() if note_col else ""
            if alias:
                entries.append(GlossaryEntry(alias=alias, canonical=canonical, note=note, source=path.name))
    return entries


def _glossary_columns(frame: pd.DataFrame) -> tuple[int, object, object, object | None]:
    alias_names = {"alias", "asr", "wrong", "term", "术语"}
    canonical_names = {"canonical", "correct", "replacement", "标准写法", "正确写法"}
    note_names = {"note", "notes", "备注"}
    first_row = [str(value).strip().lower() for value in frame.iloc[0].tolist()] if not frame.empty else []

    if any(value in alias_names | canonical_names | note_names for value in first_row):
        column_by_name = {value: frame.columns[index] for index, value in enumerate(first_row)}
        alias_col = next((column_by_name[name] for name in alias_names if name in column_by_name), frame.columns[0])
        canonical_col = next(
            (column_by_name[name] for name in canonical_names if name in column_by_name),
            frame.columns[1] if len(frame.columns) > 1 else alias_col,
        )
        note_col = next((column_by_name[name] for name in note_names if name in column_by_name), None)
        return 1, alias_col, canonical_col, note_col

    alias_col = frame.columns[0]
    canonical_col = frame.columns[1] if len(frame.columns) > 1 else alias_col
    note_col = frame.columns[2] if len(frame.columns) > 2 else None
    return 0, alias_col, canonical_col, note_col
